import sys
import os
from docx import Document

def edit_docx(filepath, comment):
    # Открываем документ
    doc = Document(filepath)

    # Добавляем текст комментария в конец документа
    if comment:
        doc.add_paragraph(comment)
    else:
        doc.add_paragraph('Комментарий отсутствует, так что просто хорошего дня)')

    # Извлекаем директорию и имя файла
    dir_name, file_name = os.path.split(filepath)
    # Добавляем "edited_" в начале имени файла
    edited_file_name = "edited_" + file_name
    # Новый путь для сохраненного файла
    edited_filepath = os.path.join(dir_name, edited_file_name)
    # Сохраняем файл с новым именем
    doc.save(edited_filepath)
    
    return edited_filepath

def main():
    if len(sys.argv) < 3:
        print("Usage: python editdocument.py <filename> <comment>")
        sys.exit(1)

    filename = sys.argv[1]
    comment = sys.argv[2]

    # Получаем абсолютный путь к директории скрипта
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, filename)

    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        sys.exit(1)

    try:
        edited_file = edit_docx(file_path, comment)
        print(f"Edited file created: {edited_file}")
    except Exception as e:
        print(f"Error editing document: {e}")
        sys.exit(1)

if __name__ == '__main__':
    main()
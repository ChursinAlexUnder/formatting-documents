import os
import sys
import tempfile
import unittest

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


SCRIPTS_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if SCRIPTS_DIR not in sys.path:
    sys.path.insert(0, SCRIPTS_DIR)

import editdocument
from modules.document_text import (
    is_table_caption,
    is_table_continuation_heading,
    normalize_caption,
    normalize_paragraph_references,
)
from modules.usednumbers import (
    findAndFormatTables,
    findBibliographyList,
    isBibliographyHeading,
)


class EditDocumentTests(unittest.TestCase):
    def test_bibliography_detection_does_not_require_ai(self):
        doc = Document()
        doc.add_paragraph("Основной текст со ссылкой [1].")
        doc.add_paragraph("СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ")
        doc.add_paragraph("1. Первый источник")
        doc.add_paragraph("2. Второй источник")

        bibliography, heading_index = findBibliographyList(doc)

        self.assertEqual(heading_index, 1)
        self.assertEqual(bibliography, [False, False])
        self.assertTrue(isBibliographyHeading(doc.paragraphs[1].text))

    def test_table_continuation_heading_is_detected(self):
        self.assertTrue(is_table_continuation_heading("Продолжение таблицы 1"))
        self.assertFalse(is_table_continuation_heading("Продолжение текста"))

    def test_annotation_source_stops_before_bibliography(self):
        doc = Document()
        doc.add_paragraph(
            "Основной содержательный текст документа с достаточным количеством "
            "слов для включения в материал аннотации."
        )
        doc.add_paragraph("СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ")
        doc.add_paragraph(
            "Секретный источник не должен отправляться модели для аннотации."
        )

        source = editdocument.build_annotation_source(doc)

        self.assertIn("Основной содержательный текст", source)
        self.assertNotIn("Секретный источник", source)

    def test_appendix_heading_and_body_use_different_rules(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            document_name = "appendix.docx"
            source_path = os.path.join(temp_dir, document_name)

            doc = Document()
            heading = doc.add_paragraph("Приложение а")
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            heading.runs[0].font.size = Pt(12)
            body = doc.add_paragraph("Текст приложения")
            body.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            body.runs[0].font.size = Pt(10)
            doc.add_paragraph("Основной текст документа.")
            doc.save(source_path)

            formatted_path, answer = editdocument.formatDocument(
                temp_dir,
                document_name,
                "Times New Roman",
                "14",
                "По ширине",
                "1.5",
                "0",
                "0",
                "1.25",
                "2.0",
                "Нет",
                None,
            )

            formatted = Document(formatted_path)
            appendix = formatted.paragraphs[0]
            appendix_body = formatted.paragraphs[1]
            self.assertEqual(appendix.text, "ПРИЛОЖЕНИЕ А")
            self.assertEqual(appendix.alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)
            self.assertTrue(appendix.runs[0].bold)
            self.assertEqual(appendix.runs[0].font.size.pt, 14)
            self.assertEqual(
                appendix_body.alignment,
                WD_PARAGRAPH_ALIGNMENT.LEFT,
            )
            self.assertEqual(appendix_body.runs[0].font.size.pt, 10)
            self.assertEqual(len(answer), 5)

    def test_reference_normalization_preserves_run_formatting(self):
        paragraph = Document().add_paragraph()
        first = paragraph.add_run("См. ")
        first.bold = True
        second = paragraph.add_run("(рис. 1)")
        second.italic = True

        normalize_paragraph_references(paragraph)

        self.assertEqual(paragraph.text, "См. (рисунок 1)")
        self.assertTrue(paragraph.runs[0].bold)
        self.assertTrue(paragraph.runs[1].italic)

    def test_numbered_heading_is_not_a_table_caption(self):
        self.assertFalse(is_table_caption("2 Экономика предприятия"))

    def test_caption_without_dash_gets_canonical_separator(self):
        paragraph = Document().add_paragraph("табл. 1 результаты")
        normalize_caption(paragraph)
        self.assertEqual(paragraph.text, "Таблица 1 – Результаты")

    def test_table_caption_is_not_counted_as_a_reference(self):
        doc = Document()
        doc.add_paragraph("Таблица 1 – Результаты")
        doc.add_table(rows=1, cols=1)
        self.assertEqual(findAndFormatTables(doc), [False])


if __name__ == "__main__":
    unittest.main()

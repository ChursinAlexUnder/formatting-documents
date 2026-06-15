import os
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from modules.annotation import generate_annotation
from modules.document_text import (
    ensure_paragraph_mark_formatting,
    is_appendix_heading,
    is_figure_caption,
    is_table_caption,
    is_table_continuation_heading,
    normalize_appendix_heading,
    normalize_caption,
    normalize_heading,
    normalize_paragraph_references,
    paragraph_will_wrap,
    replace_em_dashes,
    set_paragraph_spacing,
    trim_paragraph_text,
)
from modules.headings import (
    addEmptyParagraphAfter,
    addEmptyParagraphBefore,
    addPageBreak,
    changeNormalStyle,
    ensureHeadingStyle,
    headingLevel,
    isHeading,
)
from modules.tabs import addTab, format_list_markers, removeTabs
from modules.title import paragraphHasPageBreak
from modules.usednumbers import (
    build_figure_reference_pattern,
    collectBibliographyRegionParagraphs,
    findAndFormatTables,
    findBibliographyList,
    hasReference,
    isBibliographyHeading,
)


def format_document(bufferPath, documentName, font, fontsize, alignment,
                   spacing, beforespacing, afterspacing,
                   firstindentation, listtabulation, havetitle, openrouter_api_key=None):
    document_path = os.path.join(bufferPath, documentName)
    doc = Document(document_path)
    has_title_page = havetitle == "Есть"
    normal_size = doc.styles["Normal"].font.size
    appendix_body_formats = {}
    inside_appendix = False
    for paragraph in doc.paragraphs:
        if is_appendix_heading(paragraph.text):
            inside_appendix = True
            continue
        if isBibliographyHeading(paragraph.text):
            inside_appendix = False
        if not inside_appendix:
            continue
        fallback_size = paragraph.style.font.size or normal_size or Pt(11)
        appendix_body_formats[paragraph._element] = {
            "alignment": paragraph.alignment,
            "size": fallback_size.pt,
            "run_sizes": [
                (run.font.size or fallback_size).pt
                for run in paragraph.runs
            ],
        }
    replace_em_dashes(doc)

    haveList = False

    answer = []
    drawList = []
    drawCount = 0
    drawPattern = build_figure_reference_pattern()
    excluded_elements = set()
    title_elements = set()
    heading_spacing_requests = []

    bibliographyList, bibliographyStartIdx = findBibliographyList(doc)
    bibliographyPattern = re.compile(r"\[\s*([\d,\-–—\s]+?)\s*\]")
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    if not has_title_page:
        changeNormalStyle(doc, font, fontsize, alignment,
                          spacing, beforespacing,
                          afterspacing, firstindentation)

    isFirstPageBreak = False
    isBibliographyList = False
    body_paragraphs = list(doc.paragraphs)
    if bibliographyStartIdx is not None:
        excluded_elements.add(body_paragraphs[bibliographyStartIdx]._element)
        for paragraph in collectBibliographyRegionParagraphs(
            doc,
            bibliographyStartIdx,
        ):
            excluded_elements.add(paragraph._element)

    for abs_index, paragraph in enumerate(body_paragraphs):
        if has_title_page and not isFirstPageBreak:
            isFirstPageBreak = paragraphHasPageBreak(paragraph)
            title_elements.add(paragraph._element)
            excluded_elements.add(paragraph._element)
            continue

        isHead = False
        isDraw = False

        trim_paragraph_text(paragraph)
        normalize_paragraph_references(paragraph)
        if paragraph._element.getparent() is not None and paragraph._element.getparent().tag == qn('w:tc'):
            excluded_elements.add(paragraph._element)
        for run in paragraph.runs:
            elem = run._element
            if (elem.find(qn("w:drawing")) is not None or
                elem.find(qn("w:pict")) is not None or
                paragraph._element.xpath(".//wp:inline") or
                paragraph._element.xpath(".//wp:anchor")):
                isDraw = True
                excluded_elements.add(paragraph._element)
                drawCount += 1
                hasRef = hasReference(doc, abs_index, drawCount, drawPattern)
                drawList.append(hasRef)
                break

        figure_caption = is_figure_caption(paragraph.text)
        table_caption = is_table_caption(paragraph.text)
        continuation_heading = is_table_continuation_heading(paragraph.text)
        appendix_heading = is_appendix_heading(paragraph.text)
        appendix_body_format = appendix_body_formats.get(paragraph._element)
        if figure_caption or table_caption or continuation_heading:
            normalize_caption(paragraph)
        elif appendix_heading:
            normalize_appendix_heading(paragraph)

        heading_detected = (
            not isDraw
            and not figure_caption
            and not table_caption
            and not appendix_heading
            and isHeading(paragraph)
        )
        if heading_detected:
            normalize_heading(paragraph)

        level = headingLevel(paragraph.text)
        if paragraph._element.xpath(".//w:numPr"):
            removeTabs(paragraph)
            addTab(paragraph, listtabulation)
            if not haveList:
                haveList = True
                format_list_markers(doc, font, int(fontsize))
        elif heading_detected:
            isHead = True
            excluded_elements.add(paragraph._element)
            if level:
                style_name = ensureHeadingStyle(doc, level, font, fontsize)
                paragraph.style = style_name
                heading_spacing_requests.append((paragraph, level, abs_index))

        text_lower = paragraph.text.strip().lower()
        appendix_heading = is_appendix_heading(paragraph.text)
        figure_caption = is_figure_caption(paragraph.text)
        table_caption = is_table_caption(paragraph.text)
        continuation_heading = is_table_continuation_heading(paragraph.text)
        if isBibliographyHeading(paragraph.text):
            isBibliographyList = True

        if not isBibliographyList:
            for match in bibliographyPattern.findall(paragraph.text):
                for token in re.split(r",", match):
                    token = token.strip()
                    if any(d in token for d in ["-", "–", "—"]):
                        try:
                            start, end = map(int, re.split(r"\s*[-–—]\s*", token))
                            for i in range(start, end+1):
                                if 1 <= i <= len(bibliographyList):
                                    bibliographyList[i-1] = True
                        except ValueError:
                            pass
                    else:
                        try:
                            num = int(token)
                            if 1 <= num <= len(bibliographyList):
                                bibliographyList[num-1] = True
                        except ValueError:
                            pass
        if appendix_body_format is not None:
            paragraph.alignment = appendix_body_format["alignment"]
        elif (isDraw or figure_caption
            or text_lower in ("содержание", "введение", "заключение", "реферат")
            or appendix_heading
            or isBibliographyHeading(paragraph.text)):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            excluded_elements.add(paragraph._element)
        elif table_caption or continuation_heading:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            excluded_elements.add(paragraph._element)
        elif isHead:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif alignment == "По левому краю":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == "По центру":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == "По правому краю":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.right_indent = 0
        paragraph.paragraph_format.first_line_indent = (
            Cm(0)
            if (
                isDraw
                or figure_caption
                or table_caption
                or continuation_heading
                or text_lower in ("содержание", "введение", "заключение", "реферат")
                or appendix_heading
                or isBibliographyHeading(paragraph.text)
            )
            else Cm(float(firstindentation))
        )

        title_wraps = (
            (
                isHead
                or appendix_heading
                or figure_caption
                or table_caption
                or continuation_heading
            )
            and paragraph_will_wrap(
                doc,
                paragraph,
                font,
                fontsize,
                bold=isHead,
            )
        )
        set_paragraph_spacing(
            paragraph,
            float(fontsize) * float(beforespacing),
            float(fontsize) * float(afterspacing),
            1.0 if title_wraps else float(spacing),
        )

        paragraph_size = (
            appendix_body_format["size"]
            if appendix_body_format is not None
            else float(fontsize)
        )
        ensure_paragraph_mark_formatting(paragraph._element, font, paragraph_size)
        if appendix_body_format is None:
            paragraph.style.font.size = Pt(float(fontsize))
            paragraph.style.font.name = font

        text_runs = [run for run in paragraph.runs if run.text.strip()]
        is_code_only = bool(text_runs) and all(
            run.font.name == "Consolas" for run in text_runs
        )
        for run_index, run in enumerate(paragraph.runs):
            if run.font.name != "Consolas":
                run.font.name = font
                if appendix_body_format is not None:
                    original_sizes = appendix_body_format["run_sizes"]
                    run.font.size = Pt(
                        original_sizes[run_index]
                        if run_index < len(original_sizes)
                        else appendix_body_format["size"]
                    )
                else:
                    run.font.size = Pt(float(fontsize))
            else:
                if appendix_body_format is not None:
                    original_sizes = appendix_body_format["run_sizes"]
                    run.font.size = Pt(
                        original_sizes[run_index]
                        if run_index < len(original_sizes)
                        else appendix_body_format["size"]
                    )
                else:
                    run.font.size = Pt(11)
            if appendix_heading:
                run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        if is_code_only:
            excluded_elements.add(paragraph._element)
    for paragraph, level, abs_index in heading_spacing_requests:
        if level == 1:
            addPageBreak(paragraph)
            addEmptyParagraphAfter(paragraph)
        else:
            if abs_index > 0:
                addEmptyParagraphBefore(paragraph)
            addEmptyParagraphAfter(paragraph)
    for p_element in doc.element.body.iter(qn("w:p")):
        if p_element in title_elements:
            continue
        text = "".join(
            text_node.text or ""
            for text_node in p_element.iter(qn("w:t"))
        )
        if not text.strip():
            ensure_paragraph_mark_formatting(p_element, font, fontsize)
            set_paragraph_spacing(
                p_element,
                float(fontsize) * float(beforespacing),
                float(fontsize) * float(afterspacing),
                float(spacing),
            )
    target_half_points = int(float(fontsize) * 2)
    for hyperlink in doc.element.body.iter(qn('w:hyperlink')):
        parent = hyperlink.getparent()
        while parent is not None and parent.tag != qn("w:p"):
            parent = parent.getparent()
        if parent in appendix_body_formats:
            continue
        for r in hyperlink.iter(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            for tag in ('w:sz', 'w:szCs'):
                sz = rPr.find(qn(tag))
                if sz is None:
                    sz = OxmlElement(tag)
                    rPr.append(sz)
                sz.set(qn('w:val'), str(target_half_points))
    paragraph_count = 0
    for para in doc.paragraphs:
        if para._element in excluded_elements:
            continue
        text = para.text.strip()
        text_lower = text.lower()
        if (text_lower.startswith("рисун") or
            text_lower.startswith("таблиц") or
            text_lower.startswith("приложение") or
            isBibliographyHeading(text)):
            continue
        if re.search(r'[a-zA-Zа-яА-ЯёЁ0-9]', text):
            paragraph_count += 1

    answer.append(drawList)
    answer.append(findAndFormatTables(doc))
    answer.append(bibliographyList)
    answer.append(paragraph_count)
    annotation = ""
    if openrouter_api_key:
        annotation = generate_annotation(doc, openrouter_api_key, has_title_page)
    answer.append(annotation)

    formattedName = f"formatted_{documentName}"
    formattedPath = os.path.join(bufferPath, formattedName)
    doc.save(formattedPath)
    return formattedPath, answer

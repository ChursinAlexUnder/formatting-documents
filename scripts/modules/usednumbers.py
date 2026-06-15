import re
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from modules.document_text import (
    is_table_caption,
    is_table_continuation_heading,
    normalize_parenthetical_references,
)

ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

_BIBLIOGRAPHY_HEADING_RE = re.compile(
    r"^(?:(?:список|перечень).*(?:источник|литератур)|"
    r"библиограф(?:ический\s+список|ия))",
    re.IGNORECASE,
)
_BIBLIOGRAPHY_NUMBER_RE = re.compile(r"^\s*(\d+)[.)]\s+")


def isBibliographyHeading(text):
    return bool(_BIBLIOGRAPHY_HEADING_RE.match((text or "").strip()))


def findBibliographyHeading(doc):
    for index in range(len(doc.paragraphs) - 1, -1, -1):
        if isBibliographyHeading(doc.paragraphs[index].text):
            return index
    return None


def collectBibliographyRegionParagraphs(doc, heading_index):
    if heading_index is None:
        return []

    paragraphs = []
    found_entry = False
    for paragraph in doc.paragraphs[heading_index + 1:]:
        text = paragraph.text.strip()
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        numbered = bool(
            _BIBLIOGRAPHY_NUMBER_RE.match(text)
            or paragraph._element.xpath(".//w:numPr")
        )

        if (
            found_entry
            and text
            and (style_name.startswith("heading") or style_name.startswith("заголовок"))
            and not numbered
        ):
            break

        paragraphs.append(paragraph)
        if text and numbered:
            found_entry = True

    return paragraphs


def findBibliographyList(doc):
    """Находит список литературы и возвращает отметки ссылок для его записей."""
    bibliographyStart = findBibliographyHeading(doc)
    region = collectBibliographyRegionParagraphs(doc, bibliographyStart)
    numbered_entries = [
        paragraph
        for paragraph in region
        if paragraph.text.strip()
        and (
            _BIBLIOGRAPHY_NUMBER_RE.match(paragraph.text)
            or paragraph._element.xpath(".//w:numPr")
        )
    ]
    bibliographyLength = len(numbered_entries)
    if bibliographyLength == 0:
        bibliographyLength = sum(1 for paragraph in region if paragraph.text.strip())
    bibliographyList = [False] * bibliographyLength
    return bibliographyList, bibliographyStart

def build_figure_reference_pattern():
    return re.compile(
        r"(?:\(\s*)?(?:рис(?:ун(?:ок|ки)?|\.?|\s*-?(?:ке|ках|ку|нок))?)(?:\s+)((?:\d+(?:\s*(?:,|-|–|—)\s*\d+)*)+)(?:\s*\))?",
        re.IGNORECASE,
    )


def build_table_reference_pattern():
    return re.compile(
        r"(?:\(\s*)?(?:таб(?:лиц(?:а|ы|у|ой)?|л\.?|\.?|\s*-?(?:ца|цы|це|цу|л))?)(?:\s+)((?:\d+(?:\s*(?:,|-|–|—)\s*\d+)*)+)(?:\s*\))?",
        re.IGNORECASE,
    )


def _parse_reference_numbers(text):
    numbers = set()
    if not text:
        return numbers
    tokens = re.split(r"\s*,\s*", text)
    for token in tokens:
        token = token.strip()
        if not token:
            continue
        if any(dash in token for dash in ["-", "–", "—"]):
            try:
                dash_split = re.split(r"\s*[-–—]\s*", token)
                if len(dash_split) == 2:
                    start, end = map(int, dash_split)
                    numbers.update(range(start, end + 1))
            except ValueError:
                continue
        else:
            try:
                numbers.add(int(token))
            except ValueError:
                continue
    return numbers


def hasReference(doc, start_index, number, pattern):
    """Ищет ссылки на рисунок или таблицу в тексте, включая сокращённые формы."""
    if start_index is None:
        return False

    if start_index < 0:
        return False

    last_index = min(start_index, len(doc.paragraphs) - 1)
    for i in range(last_index, -1, -1):
        text = doc.paragraphs[i].text
        if not text:
            continue
        normalized_text = normalize_parenthetical_references(text)
        for match in pattern.findall(normalized_text):
            if number in _parse_reference_numbers(match):
                return True
    return False

def getTableParagraphIndex(doc, table):
    """
    Определяет индекс параграфа, непосредственно предшествующего таблице.
    """
    last_paragraph_index = None
    p_index = 0
    for child in doc.element.body:
        tag = child.tag
        if tag.endswith('}p'):
            last_paragraph_index = p_index
            p_index += 1
        elif tag.endswith('}tbl'):
            if child == table._element:
                break
    if last_paragraph_index is None:
        last_paragraph_index = len(doc.paragraphs) - 1
    return last_paragraph_index

def _is_continuation_table(doc, table, table_index):
    """Определяет, является ли таблица продолжением предыдущей таблицы."""
    if table_index == 0:
        return False

    previous_table = doc.tables[table_index - 1]
    previous_table_index = None
    current_table_index = None

    for child_index, child in enumerate(doc.element.body):
        if child == previous_table._element:
            previous_table_index = child_index
        elif child == table._element:
            current_table_index = child_index
            break

    if previous_table_index is None or current_table_index is None:
        return False

    meaningful_paragraphs = []
    for child in doc.element.body[previous_table_index + 1:current_table_index]:
        if child.tag.endswith('}p'):
            text = child.text.strip() if hasattr(child, 'text') else ''
            if text:
                meaningful_paragraphs.append(text)

    if not meaningful_paragraphs:
        return True

    preceding_text = " ".join(meaningful_paragraphs[-1].split())
    if is_table_continuation_heading(preceding_text):
        return True
    if is_table_caption(preceding_text):
        return False

    return False


def get_table_numbers(doc):
    """Возвращает номер логической таблицы для каждой физической таблицы в документе."""
    table_numbers = []
    current_number = 0

    for index, table in enumerate(doc.tables):
        if index == 0:
            current_number = 1
            table_numbers.append(current_number)
            continue

        if _is_continuation_table(doc, table, index):
            table_numbers.append(current_number)
        else:
            current_number += 1
            table_numbers.append(current_number)

    return table_numbers


def findAndFormatTables(doc):
    """
    Находит все таблицы в документе, форматирует их вместе с заголовками и возвращает список булевых значений,
    соответствующих наличию ссылки на таблицу.
    """
    tableList = []
    tablePattern = build_table_reference_pattern()
    table_numbers = get_table_numbers(doc)

    for index, table in enumerate(doc.tables):
        tableParagraphIndex = getTableParagraphIndex(doc, table)
        reference_end_index = tableParagraphIndex
        if 0 <= tableParagraphIndex < len(doc.paragraphs):
            heading = doc.paragraphs[tableParagraphIndex]
            if is_table_caption(heading.text):
                heading.paragraph_format.first_line_indent = Cm(0)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                reference_end_index -= 1
        hasRef = hasReference(
            doc,
            reference_end_index,
            table_numbers[index],
            tablePattern,
        )
        tableList.append(hasRef)
    return tableList

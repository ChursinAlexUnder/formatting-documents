import sys
import os
import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from modules.tabs import removeTabs, addTab
from modules.headings import headingLevel, isHeading, cycle_removeEmptyLinesAndPageBreaks, addPageBreak, addEmptyParagraphBefore, addEmptyParagraphAfter, ensureHeadingStyle, changeNormalStyle
from modules.usednumbers import findAndFormatTables, findBibliographyList, hasReference
from modules.title import paragraphHasPageBreak


def updateParagraphDefaultFont(paragraph, font):
    """Изменяет шрифт в первом <w:rFonts> внутри <w:pPr>, если он существует."""
    
    # Ищем <w:pPr> внутри параграфа
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        # Ищем <w:rPr> внутри <w:pPr>
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            # Ищем <w:rFonts> внутри <w:rPr>
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                # Меняем шрифт только если <w:rFonts> уже есть
                rFonts.set(qn("w:ascii"), font)
                rFonts.set(qn("w:hAnsi"), font)

def cleanParagraphText(paragraph):
    """Удаляет пробелы в начале и в конце текста параграфа, 
    сохраняя изображения, гиперссылки и другие элементы."""

    if not paragraph.runs:  # Если нет run'ов, ничего не делаем
        return

    # --- Удаляем пробелы в начале ---
    found_non_space = False  # Флаг, нашли ли непустой текст
    for run in paragraph.runs:
        if not found_non_space:
            if run.text.strip():  # Нашли первый run с текстом
                run.text = run.text.lstrip()  # Убираем пробелы слева
                found_non_space = True  # Дальше пробелы не трогаем
            elif run.text.isspace():  
                run.text = ""  # Полностью пробельные run'ы в начале удаляем

    # --- Удаляем пробелы в конце ---
    found_non_space = False
    for run in reversed(paragraph.runs):
        if not found_non_space:
            if run.text.strip():  # Нашли последний run с текстом
                run.text = run.text.rstrip()  # Убираем пробелы справа
                found_non_space = True  # Дальше пробелы не трогаем
            elif run.text.isspace():  
                run.text = ""  # Полностью пробельные run'ы в конце удаляем

def modifyList(doc, font, fontsize):
    """
    Изменяет стиль номеров или маркеров списка в документе.
    """
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part.element

    for abstract_num in numbering_xml.findall(qn("w:abstractNum")):
        for lvl in abstract_num.findall(qn("w:lvl")):
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None:
                continue

            num_fmt_val = num_fmt.get(qn("w:val"))

            # Настройки шрифта
            rPr = lvl.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                lvl.append(rPr)

            if num_fmt_val != "bullet": #нумерованный список
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                rFonts.set(qn("w:ascii"), font)
                rFonts.set(qn("w:hAnsi"), font)

            sz = rPr.find(qn("w:sz"))
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(qn("w:val"), str(fontsize * 2))

def formatDocument(bufferPath, documentName, font, fontsize, alignment,
                   spacing, beforespacing, afterspacing,
                   firstindentation, listtabulation, havetitle):
    # Открываем документ
    doc = Document(f"{bufferPath}/{documentName}")

    haveList = False
    isDrawTitle = False

    answer = []
    drawList = []
    drawCount = 0
    drawPattern = re.compile(
        r"(?:\(\s*)?рисун\w*\s+((?:\d+(?:\s*[-,–—]\s*)?)+)(?:\s*\))?",
        re.IGNORECASE
    )

    bibliographyList = findBibliographyList(doc)
    bibliographyPattern = re.compile(r"\[\s*([\d,\-–—\s]+?)\s*\]")

    # Настройка полей для основного раздела
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Настройка стиля Normal, если нет титульного листа
    if havetitle == "Нет":
        changeNormalStyle(doc, font, fontsize, alignment,
                          spacing, beforespacing,
                          afterspacing, firstindentation)

    isFirstPageBreak = False
    isBibliographyList = False

    # Основной цикл: перебираем все абзацы с помощью enumerate
    for abs_index, paragraph in enumerate(doc.paragraphs):
        # Пропускаем титульный лист, но не теряем синхронизацию индекса
        if havetitle == "Есть" and not isFirstPageBreak:
            isFirstPageBreak = paragraphHasPageBreak(paragraph)
            if isFirstPageBreak:
                havetitle = "Нет"
            continue

        isHead = False
        isDraw = False

        cleanParagraphText(paragraph)
        level = headingLevel(paragraph.text)

        # Поиск рисунка: XML-элементы w:drawing, w:pict, wp:inline, wp:anchor
        for run in paragraph.runs:
            elem = run._element
            if (elem.find(qn("w:drawing")) is not None or
                elem.find(qn("w:pict")) is not None or
                paragraph._element.xpath(".//wp:inline") or
                paragraph._element.xpath(".//wp:anchor")):
                isDraw = True
                isDrawTitle = True
                drawCount += 1
                # Проверяем наличие ссылки на рисунок до этой позиции
                hasRef = hasReference(doc, abs_index, drawCount, drawPattern)
                drawList.append(hasRef)
                break

        # Обработка списков
        if paragraph._element.xpath(".//w:numPr"):
            removeTabs(paragraph)
            addTab(paragraph, listtabulation)
            if not haveList:
                haveList = True
                modifyList(doc, font, int(fontsize))

        # Обработка заголовков
        elif not isDraw and not isDrawTitle and isHeading(paragraph) and level:
            style_name = ensureHeadingStyle(doc, level, font, fontsize)
            paragraph.style = style_name
            isHead = True

            removed = cycle_removeEmptyLinesAndPageBreaks(doc)
            # нет нужды корректировать abs_index вручную

            if level == 1:
                addPageBreak(paragraph)
                addEmptyParagraphAfter(paragraph)
            else:
                if abs_index > 0:
                    addEmptyParagraphBefore(paragraph)
                addEmptyParagraphAfter(paragraph)

        # Убираем все элементы w:spacing у абзаца
        for spacing_elem in paragraph._element.xpath('.//w:spacing'):
            spacing_elem.getparent().remove(spacing_elem)

        text_lower = paragraph.text.strip().lower()
        if text_lower.startswith("список") and ("источников" in text_lower or "литературы" in text_lower):
            isBibliographyList = True

        # Поиск ссылок на литературу до секции списка
        if not isBibliographyList:
            for match in bibliographyPattern.findall(paragraph.text):
                for token in re.split(r",", match):
                    token = token.strip()
                    # диапазоны и одиночные номера
                    if any(d in token for d in ["-", "–", "—"]):
                        try:
                            start, end = map(int, re.split(r"\s*[-–—]\s*", token))
                            for i in range(start, end+1):
                                if 1 <= i <= len(bibliographyList):
                                    bibliographyList[i-1] = True
                        except ValueError:
                            pass
                    else:
                        try:
                            num = int(token)
                            if 1 <= num <= len(bibliographyList):
                                bibliographyList[num-1] = True
                        except ValueError:
                            pass

        # Выравнивание и форматирование абзаца
        if (isDraw or (isDrawTitle and text_lower.startswith("рисун"))
            or text_lower in ("содержание", "введение", "заключение", "реферат", "приложение")
            or (text_lower.startswith("список") and ("источников" in text_lower or "литературы" in text_lower))):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif isHead:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif alignment == "По левому краю":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == "По центру":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == "По правому краю":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        paragraph.paragraph_format.line_spacing = float(spacing)
        paragraph.paragraph_format.space_before = Pt(float(fontsize) * float(beforespacing))
        paragraph.paragraph_format.space_after = Pt(float(fontsize) * float(afterspacing))
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.right_indent = 0
        paragraph.paragraph_format.first_line_indent = Cm(0) if isDraw or isDrawTitle or text_lower in ("содержание", "введение", "заключение", "реферат", "приложение") or (text_lower.startswith("список") and ("источников" in text_lower or "литературы" in text_lower)) else Cm(float(firstindentation))

        updateParagraphDefaultFont(paragraph, font)
        paragraph.style.font.size = Pt(float(fontsize))
        paragraph.style.font.name = font

        for run in paragraph.runs:
            if run.font.name != "Consolas":
                run.font.name = font
                run.font.size = Pt(float(fontsize))
            else:
                run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Сброс флага заголовка рисунка после первого текстового абзаца
        if isDrawTitle and not isDraw:
            isDrawTitle = False

    # Добавление списка рисунков и таблиц, библиографии
    answer.append(drawList)
    answer.append(findAndFormatTables(doc))
    answer.append(bibliographyList)

    # Сохранение
    formattedName = f"formatted_{documentName}"
    formattedPath = f"{bufferPath}/{formattedName}"
    doc.save(formattedPath)
    return formattedPath, answer

documentName = sys.argv[1]
bufferPath = '../buffer'
documentPath = bufferPath + '/' + documentName
font = sys.argv[2]
fontsize = sys.argv[3]
alignment = sys.argv[4]
spacing = sys.argv[5]
beforespacing = sys.argv[6]
afterspacing = sys.argv[7]
firstindentation = sys.argv[8]
listtabulation = sys.argv[9]
havetitle = sys.argv[10]

if not os.path.exists(documentPath):
    print(f"Document not found: {documentPath}")
    sys.exit(1)

try:
    formattedDocumentPath, result = formatDocument(bufferPath, documentName, font, fontsize, alignment, spacing, beforespacing, afterspacing, firstindentation, listtabulation, havetitle)
    print(json.dumps(result))
except Exception as e:
    print(f"Error formatting document: {e}")
    sys.exit(1)